import subprocess
import os
import sys
import docx
import re

python_bin = r"C:\salo\acb\quill\env_quill\Scripts\python.exe"
convert_script = "convert_md.py"

print("==================================================")
print("Verifying Specifications & Work in plan.md")
print("==================================================\n")

# Create a test markdown file with various features to be verified
test_md_content = """---
title: "Plan Spec Verification"
author: "Verify Bot"
lang: "en-US"
---

# Heading 1
### Heading 3 (Skipped level 2 - should trigger warning)

Here is a paragraph with a link to [click here](http://example.com) (non-descriptive - should trigger warning).

Here is an image: ![](inaccessible_documents/inaccessible_image1.png) (missing alt text - should be fixed to user notice).

### Equations
Inline math: t:{a^2 + b^2 = c^2}

Block math:
t:{f(x) = sum_(i=0)^n x_i}

### Tables
| Header A | | Header C |
| --- | --- | --- |
| Cell 1 | .. | Cell 3 |
| Cell 4 | Cell 5 | Cell 6 |
"""

test_md_file = "spec_test.md"
with open(test_md_file, "w", encoding="utf-8") as f:
    f.write(test_md_content)

print(f"Created temporary spec test file: {test_md_file}")

results = {}

# --- Stage 1 & 5: Markdown validation, alt-text fix, and console output check ---
print("\n--- Testing Stage 1 (Warnings & Alt Text) and Stage 5 (Axe/Tidy Validation) ---")
cmd = [python_bin, convert_script, test_md_file, "html"]
res = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")

output_stdout = res.stdout or ""
output_stderr = res.stderr or ""
full_output = output_stdout + "\n" + output_stderr

# Verify warning messages
has_heading_warning = "Heading level skipped" in full_output
has_link_warning = "non-descriptive text" in full_output
has_image_warning = "missing a descriptive alt text" in full_output
has_axe_audit = "Axe CLI accessibility audit" in full_output or "Warning: Axe CLI" in full_output or "Axe Accessibility Audit" in full_output

results["Stage 1: Heading hierarchy skipped check"] = "PASSED" if has_heading_warning else "FAILED"
results["Stage 1: Non-descriptive link label check"] = "PASSED" if has_link_warning else "FAILED"
results["Stage 1: Image alt text audit warning check"] = "PASSED" if has_image_warning else "FAILED"
results["Stage 5: Axe accessibility auditor invocation"] = "PASSED" if has_axe_audit else "FAILED"

# --- Stage 6: Output file default naming check ---
expected_html_out = "spec_test_out.html"
has_default_naming = os.path.exists(expected_html_out)
results["Stage 6: Output default naming (_out appended)"] = "PASSED" if has_default_naming else "FAILED"

if has_default_naming:
    # Read the output HTML to verify Stage 2 features
    print("\n--- Testing Stage 2 (HTML WCAG 2.2 Compliance) ---")
    with open(expected_html_out, "r", encoding="utf-8") as f:
        html_content = f.read()
        
    has_lang = 'lang="en-US"' in html_content
    has_title = '<title>Plan Spec Verification</title>' in html_content
    has_skip_link = 'Skip to main content' in html_content and 'href="#main-content"' in html_content
    has_main_tag = '<main id="main-content">' in html_content and '</main>' in html_content
    has_acc_css = 'line-height: 1.65' in html_content and 'letter-spacing:' in html_content
    has_focus_visible = ':focus-visible' in html_content
    has_mathml_fallback = 'alttext="Formula:' in html_content
    has_math_container = 'class="math-container"' in html_content
    
    # Check alt-text fixer dynamically (look for regex match)
    alt_fix_match = re.search(r'alt="this image in this \d+ is not alt text \. put it"', html_content)
    has_alt_fix = alt_fix_match is not None
    if has_alt_fix:
        print(f"    Dynamic alt-text fixer matched: {alt_fix_match.group(0)}")
        
    has_colspan = 'colspan="2"' in html_content
    
    results["Stage 2: Language attribute (lang=\"en-US\") check"] = "PASSED" if has_lang else "FAILED"
    results["Stage 2: Standard Title tag check"] = "PASSED" if has_title else "FAILED"
    results["Stage 2: Skip-to-content accessibility link"] = "PASSED" if has_skip_link else "FAILED"
    results["Stage 2: Wrap body inside main landmarks"] = "PASSED" if has_main_tag else "FAILED"
    results["Stage 2: Accessibility high contrast CSS injected"] = "PASSED" if has_acc_css else "FAILED"
    results["Stage 2: Visual keyboard focus visible outline check"] = "PASSED" if has_focus_visible else "FAILED"
    results["Stage 2: MathML enriched with LaTeX alttext fallback"] = "PASSED" if has_mathml_fallback else "FAILED"
    results["Stage 2: Display math wrapped in focusable scroll container"] = "PASSED" if has_math_container else "FAILED"
    results["Stage 2: Missing image alt text remediation replacement"] = "PASSED" if has_alt_fix else "FAILED"
    results["Stage 2: Empty/merge cell column index span (colspan)"] = "PASSED" if has_colspan else "FAILED"
    
    # Clean up HTML file
    try:
        os.remove(expected_html_out)
    except:
        pass

# --- Stage 3: Word DOCX Post-processing check ---
print("\n--- Testing Stage 3 (Word DOCX Post-Processing) ---")
expected_docx_out = "spec_test_out.docx"
cmd = [python_bin, convert_script, test_md_file, "word"]
res = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")

if os.path.exists(expected_docx_out):
    doc = docx.Document(expected_docx_out)
    
    has_tbl_header = False
    has_cant_split = False
    has_cell_merging = False
    
    if doc.tables:
        table = doc.tables[0]
        has_tbl_header = bool(table.rows[0]._tr.xpath('w:trPr/w:tblHeader'))
        has_cant_split = all(bool(row._tr.xpath('w:trPr/w:cantSplit')) for row in table.rows)
        # Compare cell underlying XML element pointers (_tc) to check merging
        has_cell_merging = table.rows[1].cells[0]._tc == table.rows[1].cells[1]._tc
        print(f"    Table check: Row 1 Cell 0 _tc = {table.rows[1].cells[0]._tc}, Cell 1 _tc = {table.rows[1].cells[1]._tc}, equal = {has_cell_merging}")
        
    results["Stage 3: MS Word table repeating header row (tblHeader)"] = "PASSED" if has_tbl_header else "FAILED"
    results["Stage 3: MS Word row split prevention (cantSplit)"] = "PASSED" if has_cant_split else "FAILED"
    results["Stage 3: MS Word table cell merging (colspan mapping)"] = "PASSED" if has_cell_merging else "FAILED"
    
    try:
        os.remove(expected_docx_out)
    except:
        pass
else:
    results["Stage 3: MS Word table repeating header row (tblHeader)"] = "FAILED"
    results["Stage 3: MS Word row split prevention (cantSplit)"] = "FAILED"
    results["Stage 3: MS Word table cell merging (colspan mapping)"] = "FAILED"

# --- Stage 6: Self-conversion check ---
print("\n--- Testing Stage 6 (Self-Conversions) ---")
dummy_docx = "dummy_src.docx"
d = docx.Document()
d.add_paragraph("This is a word self-conversion test.")
d.save(dummy_docx)

cmd = [python_bin, convert_script, dummy_docx, "word"]
res = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")

expected_self_docx = "dummy_src_out.docx"
has_self_docx = os.path.exists(expected_self_docx)
results["Stage 6: Word-to-Word Self-Conversion support"] = "PASSED" if has_self_docx else "FAILED"

# Cleanup
for f_path in (test_md_file, dummy_docx, expected_self_docx):
    if os.path.exists(f_path):
        try:
            os.remove(f_path)
        except:
            pass

# Print final report table
print("\n" + "="*80)
print(f"{'SPECIFICATION CHECK / ACTION ITEM':<60} | {'STATUS':<15}")
print("-"*80)
for spec, status in results.items():
    print(f"{spec:<60} | {status:<15}")
print("="*80 + "\n")

all_passed = all(status == "PASSED" for status in results.values())
if all_passed:
    print("All specifications and plan milestones are fully verified!")
    sys.exit(0)
else:
    print("Some specification verification checks failed!")
    sys.exit(1)
